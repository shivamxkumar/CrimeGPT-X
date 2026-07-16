"""Document export — real HTML → PDF / DOCX rendering of generated document content."""
import io
import logging
from bs4 import BeautifulSoup

from app.worker import celery_app

logger = logging.getLogger(__name__)


def render_pdf(html_content: str) -> bytes:
    """Render real document HTML to PDF bytes via WeasyPrint."""
    from weasyprint import HTML
    return HTML(string=html_content).write_pdf()


def render_docx(title: str, html_content: str) -> bytes:
    """Render real document HTML to a DOCX file via python-docx, preserving
    headings/paragraphs/tables extracted from the actual generated content."""
    from docx import Document as DocxDocument

    soup = BeautifulSoup(html_content, "html.parser")
    doc = DocxDocument()
    doc.add_heading(title, level=0)

    body = soup.body or soup
    for el in body.find_all(["h1", "h2", "h3", "p", "li", "table"], recursive=True):
        if el.name in ("h1", "h2", "h3"):
            level = int(el.name[1])
            text = el.get_text(strip=True)
            if text:
                doc.add_heading(text, level=level)
        elif el.name == "table":
            rows = el.find_all("tr")
            if not rows:
                continue
            n_cols = max(len(r.find_all(["td", "th"])) for r in rows)
            table = doc.add_table(rows=0, cols=n_cols)
            for r in rows:
                cells = r.find_all(["td", "th"])
                row_cells = table.add_row().cells
                for i, cell in enumerate(cells):
                    if i < n_cols:
                        row_cells[i].text = cell.get_text(strip=True)
        elif el.name in ("p", "li"):
            text = el.get_text(strip=True)
            if text:
                doc.add_paragraph(text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@celery_app.task(name="app.tasks.doc_tasks.export_pdf")
def export_pdf(document_id: str, html_content: str):
    """Background PDF export of a generated document's real content."""
    try:
        pdf_bytes = render_pdf(html_content)
        logger.info(f"PDF exported for document {document_id} ({len(pdf_bytes)} bytes)")
        return {"document_id": document_id, "size": len(pdf_bytes)}
    except Exception as e:
        logger.error(f"PDF export failed for document {document_id}: {e}")
        return {"error": str(e)}


@celery_app.task(name="app.tasks.doc_tasks.export_docx")
def export_docx(document_id: str, title: str, html_content: str):
    """Background DOCX export of a generated document's real content."""
    try:
        docx_bytes = render_docx(title, html_content)
        logger.info(f"DOCX exported for document {document_id} ({len(docx_bytes)} bytes)")
        return {"document_id": document_id, "size": len(docx_bytes)}
    except Exception as e:
        logger.error(f"DOCX export failed for document {document_id}: {e}")
        return {"error": str(e)}
